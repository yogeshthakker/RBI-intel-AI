"""
Anonymization dictionary layer for rbi-intel.
Standalone first — validates term <-> placeholder round-trip
before it gets wired into scaffold.py.

DB: creates table `anon_dictionary` in the same SQLite file rbi-intel uses.
Default path matches rbi-intel's default: ~/.rbi-intel/regdata.db
Override with RBI_INTEL_DB env var (same variable rbi-intel already uses).

Usage:
  python anon_dictionary.py add --term "Sber India" --category bank_name
  python anon_dictionary.py add --term "Mumbai Fort Branch" --category branch_name
  python anon_dictionary.py list
  python anon_dictionary.py anonymize --file policy.txt
  python anon_dictionary.py deanonymize --file scaffold_output.txt
  python anon_dictionary.py roundtrip --file policy.txt   # sanity test
"""
import argparse
import os
import re
import sqlite3
import sys
from pathlib import Path

def get_db_path():
    env = os.environ.get("RBI_INTEL_DB")
    if env:
        return env
    return str(Path.home() / ".rbi-intel" / "regdata.db")

def get_conn():
    db_path = get_db_path()
    if not Path(db_path).exists():
        print(f"!! DB not found at {db_path}. Set RBI_INTEL_DB or check rbi-intel setup.")
        sys.exit(1)
    conn = sqlite3.connect(db_path)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS anon_dictionary (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            term TEXT NOT NULL UNIQUE,
            placeholder TEXT NOT NULL UNIQUE,
            category TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    return conn

def next_placeholder(conn, category):
    cat_tag = re.sub(r"[^A-Z]", "", category.upper())[:4] or "TERM"
    row = conn.execute(
        "SELECT COUNT(*) FROM anon_dictionary WHERE category = ?", (category,)
    ).fetchone()
    n = row[0] + 1
    return f"[{cat_tag}_{n:03d}]"

def add_term(term, category, placeholder=None):
    conn = get_conn()
    if placeholder is None:
        placeholder = next_placeholder(conn, category)
    try:
        conn.execute(
            "INSERT INTO anon_dictionary (term, placeholder, category) VALUES (?, ?, ?)",
            (term, placeholder, category),
        )
        conn.commit()
        print(f"added: {term!r} -> {placeholder}")
    except sqlite3.IntegrityError as e:
        print(f"!! skipped (already exists?): {term!r} — {e}")
    conn.close()

def list_terms():
    conn = get_conn()
    rows = conn.execute(
        "SELECT term, placeholder, category FROM anon_dictionary ORDER BY category, term"
    ).fetchall()
    conn.close()
    if not rows:
        print("(dictionary is empty — add terms first)")
        return
    for term, placeholder, category in rows:
        print(f"[{category}]  {term!r} <-> {placeholder}")

def _load_map(conn):
    return conn.execute("SELECT term, placeholder FROM anon_dictionary").fetchall()

def anonymize_text(text):
    conn = get_conn()
    pairs = _load_map(conn)
    conn.close()
    # longest term first so substrings don't clobber longer matches
    pairs.sort(key=lambda p: len(p[0]), reverse=True)
    for term, placeholder in pairs:
        text = re.sub(re.escape(term), placeholder, text, flags=re.IGNORECASE)
    return text

def deanonymize_text(text):
    conn = get_conn()
    pairs = _load_map(conn)
    conn.close()
    for term, placeholder in pairs:
        text = text.replace(placeholder, term)
    return text

def read_any_text(file_path):
    """Reads .txt directly. Reads .pdf by extracting text via pypdf."""
    p = Path(file_path)
    if p.suffix.lower() == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            print("!! pypdf not installed. Run: pip install pypdf --break-system-packages")
            print("   (on your machine, just: pip install pypdf)")
            sys.exit(1)
        reader = PdfReader(str(p))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages)
        if not text.strip():
            print("!! Warning: extracted 0 characters — this PDF may be scanned/image-based, needs OCR.")
        return text
    if p.suffix.lower() == ".docx":
        try:
            import docx
        except ImportError:
            print("!! python-docx not installed. Run: pip install python-docx --break-system-packages")
            print("   (on your machine, just: pip install python-docx)")
            sys.exit(1)
        doc = docx.Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts)
        if not text.strip():
            print("!! Warning: extracted 0 characters from this .docx.")
        return text
    return p.read_text(encoding="utf-8")

def cmd_anonymize(file_path):
    text = read_any_text(file_path)
    out = anonymize_text(text)
    out_path = Path(file_path).with_suffix(".anon.txt")
    out_path.write_text(out, encoding="utf-8")
    print(f"wrote {out_path}")

def cmd_deanonymize(file_path):
    text = read_any_text(file_path)
    out = deanonymize_text(text)
    out_path = Path(file_path).with_suffix(".deanon.txt")
    out_path.write_text(out, encoding="utf-8")
    print(f"wrote {out_path}")

def cmd_roundtrip(file_path):
    original = read_any_text(file_path)
    anon = anonymize_text(original)
    back = deanonymize_text(anon)
    print("--- ANONYMIZED (first 500 chars) ---")
    print(anon[:500])
    print("\n--- ROUND-TRIP CHECK ---")
    if back == original:
        print("PASS: de-anonymized text matches original exactly.")
    else:
        print("FAIL: mismatch after round trip.")
        for i, (a, b) in enumerate(zip(original, back)):
            if a != b:
                print(f"  first diff at char {i}: original={a!r} back={b!r}")
                break
        print(f"  len original={len(original)} len back={len(back)}")

SPACY_LABEL_TO_CATEGORY = {
    "ORG": "org_name",
    "PERSON": "person_name",
    "GPE": "location",     # countries, cities, states
    "LOC": "location",
    "FAC": "facility_name",  # buildings, branches, named facilities
}

def cmd_scan(file_path, min_count):
    try:
        import spacy
    except ImportError:
        print("!! spaCy not installed. Run:")
        print("   pip install spacy")
        print("   python -m spacy download en_core_web_sm")
        sys.exit(1)
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        print("!! spaCy model not downloaded. Run:")
        print("   python -m spacy download en_core_web_sm")
        sys.exit(1)

    text = read_any_text(file_path)
    if not text.strip():
        print("!! No text extracted from file — nothing to scan.")
        return

    conn = get_conn()
    existing_terms = {t.lower() for t, _ in _load_map(conn)}
    conn.close()

    # spaCy has an internal doc-length cap; chunk long policy docs so nothing silently truncates.
    CHUNK = 100_000
    counts = {}  # (text_lower, label) -> {"text": original_case, "count": n}
    for start in range(0, len(text), CHUNK):
        piece = text[start:start + CHUNK]
        doc = nlp(piece)
        for ent in doc.ents:
            if ent.label_ not in SPACY_LABEL_TO_CATEGORY:
                continue
            cleaned = ent.text.strip()
            if len(cleaned) < 3 or cleaned.lower() in existing_terms:
                continue
            key = (cleaned.lower(), ent.label_)
            if key not in counts:
                counts[key] = {"text": cleaned, "count": 0}
            counts[key]["count"] += 1

    results = [v for v in counts.values() if v["count"] >= min_count]
    results.sort(key=lambda v: -v["count"])

    if not results:
        print("(no new candidate terms found above --min-count threshold)")
        return

    print(f"Found {len(results)} candidate identifying terms not yet in your dictionary:\n")
    by_label = {}
    for (text_lower, label), v in counts.items():
        if v["count"] < min_count:
            continue
        by_label.setdefault(label, []).append(v)

    for label, items in by_label.items():
        category = SPACY_LABEL_TO_CATEGORY[label]
        print(f"--- {label} -> suggested category: {category} ---")
        for v in sorted(items, key=lambda x: -x["count"]):
            print(f"  {v['count']:>3}x  {v['text']!r}")
        print()

    print("Review the list above. To add one, run:")
    print('  python anon_dictionary.py add --term "EXACT TEXT" --category CATEGORY')
    print("Nothing was added automatically — this only suggests, you approve.")

def main():
    p = argparse.ArgumentParser()
    sub = p.add_subparsers(dest="cmd", required=True)

    a = sub.add_parser("add")
    a.add_argument("--term", required=True)
    a.add_argument("--category", required=True, help="e.g. bank_name, branch_name, person_name, system_name")
    a.add_argument("--placeholder", default=None, help="optional, auto-generated if omitted")

    sub.add_parser("list")

    az = sub.add_parser("anonymize")
    az.add_argument("--file", required=True)

    dz = sub.add_parser("deanonymize")
    dz.add_argument("--file", required=True)

    rt = sub.add_parser("roundtrip")
    rt.add_argument("--file", required=True)

    sc = sub.add_parser("scan")
    sc.add_argument("--file", required=True)
    sc.add_argument("--min-count", type=int, default=1, help="only suggest terms appearing at least this many times")

    args = p.parse_args()

    if args.cmd == "add":
        add_term(args.term, args.category, args.placeholder)
    elif args.cmd == "list":
        list_terms()
    elif args.cmd == "anonymize":
        cmd_anonymize(args.file)
    elif args.cmd == "deanonymize":
        cmd_deanonymize(args.file)
    elif args.cmd == "roundtrip":
        cmd_roundtrip(args.file)
    elif args.cmd == "scan":
        cmd_scan(args.file, args.min_count)

if __name__ == "__main__":
    main()
